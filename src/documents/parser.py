import pdfplumber
import docx
import chardet
import os
import io
from typing import List, Union
from src.observability.logger import get_logger

logger = get_logger(__name__)

class DocumentParser:
    def parse(self, file_path: str, filename: str) -> List[dict]:
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf':
            return self._parse_pdf(file_path, filename)
        if ext == '.docx':
            return self._parse_docx(file_path, filename)
        if ext in ('.txt', '.md'):
            return self._parse_txt(file_path, filename)
        raise ValueError('Unsupported file type')

    def parse_file(self, uploaded_file) -> str:
        """Parse a Streamlit UploadedFile and return concatenated text content"""
        filename = uploaded_file.name
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            if ext == '.pdf':
                return self._parse_pdf_bytes(uploaded_file.getbuffer(), filename)
            elif ext == '.docx':
                return self._parse_docx_bytes(uploaded_file.getbuffer(), filename)
            elif ext in ('.txt', '.md'):
                return self._parse_txt_bytes(uploaded_file.getbuffer(), filename)
            else:
                raise ValueError(f'Unsupported file type: {ext}')
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {str(e)}")
            raise

    def _parse_pdf_bytes(self, file_bytes, filename: str) -> str:
        """Parse PDF from bytes buffer"""
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ''
                if text:
                    text_parts.append(text)
        return '\n\n'.join(text_parts)

    def _parse_docx_bytes(self, file_bytes, filename: str) -> str:
        """Parse DOCX from bytes buffer"""
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        return '\n\n'.join(paragraphs)

    def _parse_txt_bytes(self, file_bytes, filename: str) -> str:
        """Parse TXT from bytes buffer"""
        enc = chardet.detect(file_bytes)['encoding'] or 'utf-8'
        text = file_bytes.decode(enc, errors='replace')
        return text

    def _chunk_words(self, text, chunk_tokens=512, overlap=100):
        words = text.split()
        token_per_word = 0.75
        chunk_words = int(chunk_tokens * token_per_word)
        overlap_words = int(overlap * token_per_word)
        chunks = []
        i = 0
        idx = 0
        while i < len(words):
            part = words[i:i+chunk_words]
            chunks.append((idx, ' '.join(part)))
            idx += 1
            i += max(1, chunk_words-overlap_words)
        return chunks

    def _parse_pdf(self, path: str, filename: str) -> List[dict]:
        out = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ''
                chunks = self._chunk_words(text)
                for idx, c in chunks:
                    out.append({'text': c, 'filename': filename, 'page_number': i, 'chunk_index': idx, 'char_count': len(c), 'token_estimate': int(len(c.split())/0.75)})
        return out

    def _parse_docx(self, path: str, filename: str) -> List[dict]:
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        out = []
        text = '\n'.join(paragraphs)
        chunks = self._chunk_words(text)
        for idx, c in chunks:
            out.append({'text': c, 'filename': filename, 'page_number': int(idx/30)+1, 'chunk_index': idx, 'char_count': len(c), 'token_estimate': int(len(c.split())/0.75)})
        return out

    def _parse_txt(self, path: str, filename: str) -> List[dict]:
        with open(path, 'rb') as f:
            b = f.read()
        enc = chardet.detect(b)['encoding'] or 'utf-8'
        text = b.decode(enc, errors='replace')
        chunks = self._chunk_words(text)
        out = []
        for idx, c in chunks:
            out.append({'text': c, 'filename': filename, 'page_number': None, 'chunk_index': idx, 'char_count': len(c), 'token_estimate': int(len(c.split())/0.75)})
        return out

if __name__ == '__main__':
    p = DocumentParser()
    print('Parser ready')
